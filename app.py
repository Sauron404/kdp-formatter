import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os

try:
    from docx2pdf import convert as docx_to_pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# Crea una pagina centrata con un certo testo
def add_centered_page(doc, lines):
    doc.add_page_break()
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(16)
        run.font.name = 'Georgia'

# Aggiungi numeri di pagina nel piè di pagina
def add_page_numbers(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Aggiunge l'indice automatico con i livelli Heading
def add_table_of_contents(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = r'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Verifica se uno stile esiste nel documento
def style_exists(doc, style_name):
    try:
        _ = doc.styles[style_name]
        return True
    except KeyError:
        return False

# Verifica se un paragrafo può essere un titolo
def is_probable_title(paragraph):
    text = paragraph.text.strip()
    if not text or len(text) > 80:
        return False
    is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
    is_upper = text.isupper()
    has_larger_font = any(run.font.size and run.font.size.pt >= 14 for run in paragraph.runs)
    conditions_met = sum([is_bold, is_upper, has_larger_font])
    return conditions_met >= 2

# Formatta il documento
def format_docx(uploaded_file, formato="cartaceo", frontespizio=True, numeri_pagina=True, titolo_libro="Titolo del Libro", autore_libro="Autore", editore="Nome Editore"):
    original_doc = Document(uploaded_file)
    doc = Document()

    if frontespizio:
        add_centered_page(doc, [titolo_libro, autore_libro])
        add_centered_page(doc, [f"© 2025 {editore}", "Tutti i diritti riservati"])
        doc.add_page_break()
        doc.add_paragraph("Indice")
        add_table_of_contents(doc)
        doc.add_page_break()

    has_heading1 = style_exists(doc, 'Heading 1')

    for para in original_doc.paragraphs:
        if is_probable_title(para):
            doc.add_page_break()
            new_p = doc.add_paragraph(para.text)
            if has_heading1:
                new_p.style = 'Heading 1'
        else:
            new_p = doc.add_paragraph(para.text)

        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in new_p.runs:
            run.font.name = 'Georgia'
            run.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.87)
        section.right_margin = Inches(0.67)
        if numeri_pagina and formato == "cartaceo":
            add_page_numbers(section)

    return doc

# Streamlit UI
st.set_page_config(page_title="KDP Formatter", layout="centered")
st.title("\U0001F4D8 KDP Formatter 6x9")
st.write("Carica un file Word `.docx` e ottieni un file formattato per la stampa o l'eBook.")

uploaded_file = st.file_uploader("\U0001F4E4 Carica il tuo file Word", type=["docx"])
formato = st.selectbox("\U0001F58B️ Formato desiderato:", ["cartaceo", "ebook"])
add_frontespizio = st.checkbox("Aggiungi frontespizio?", value=True)
add_numeri_pagina = st.checkbox("Aggiungi numeri di pagina?", value=True)
titolo_libro = st.text_input("Titolo del libro:", "Titolo del Libro")
autore_libro = st.text_input("Autore:", "Autore")
editore = st.text_input("Editore:", "Nome Editore")

if uploaded_file:
    if st.button("\U0001F4C4 Formatta il documento"):
        with st.spinner("Formattazione in corso..."):
            doc = format_docx(
                uploaded_file,
                formato=formato,
                frontespizio=add_frontespizio,
                numeri_pagina=add_numeri_pagina,
                titolo_libro=titolo_libro,
                autore_libro=autore_libro,
                editore=editore
            )

            filename_base = titolo_libro.strip().lower().replace(" ", "_") or "kdp_formattato"

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                docx_path = tmp.name

            with open(docx_path, "rb") as f:
                st.success("✅ Documento Word formattato con successo!")
                st.download_button(
                    label="\U0001F4E5 Scarica .DOCX",
                    data=f,
                    file_name=f"{filename_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            if DOCX2PDF_AVAILABLE:
                with tempfile.TemporaryDirectory() as tmpdir:
                    pdf_path = os.path.join(tmpdir, "output.pdf")
                    docx_to_pdf(docx_path, pdf_path)
                    with open(pdf_path, "rb") as pdf_file:
                        st.download_button(
                            label="\U0001F4C4 Scarica anche in PDF",
                            data=pdf_file,
                            file_name=f"{filename_base}.pdf",
                            mime="application/pdf"
                        )

            os.remove(docx_path)
